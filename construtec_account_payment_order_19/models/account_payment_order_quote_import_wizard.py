import base64
import logging

from odoo import api, fields, models
from odoo.exceptions import UserError

from ..tools.quote_extraction_api import (
    QuoteExtractionError, extract_quote_from_image, extract_quote_from_pdf, extract_quote_from_text,
)

_logger = logging.getLogger(__name__)

_IMAGE_MIME_BY_EXT = {
    '.png': 'image/png',
    '.jpg': 'image/jpeg',
    '.jpeg': 'image/jpeg',
    '.webp': 'image/webp',
}


def _normalizar(texto):
    """Normaliza texto para comparar "¿ya existe?" de forma simple y predecible - mayúsculas,
    sin espacios de más. Comparación EXACTA tras normalizar, nunca asistida por IA ni difusa
    (decisión explícita del usuario) - un nombre escrito distinto no calza y se crea una
    entrada nueva en el catálogo, en vez de adivinar cuál es "la misma" cosa."""
    return ' '.join((texto or '').strip().upper().split())


def _extraer_texto_docx(docx_bytes):
    """Extrae texto de un .docx - párrafos Y tablas (una cotización casi siempre viene en
    tabla: material/cantidad/precio), concatenados en un solo texto para pasarlo a la IA como
    texto plano. Import perezoso (mismo patrón que pdfminer.six en
    construtec_account_19/models/sat_retention_import.py) - así este módulo sigue cargando en
    cualquier entorno que todavía no tenga python-docx instalado; solo revienta si alguien
    realmente sube un .docx sin la librería presente."""
    import io
    from docx import Document

    doc = Document(io.BytesIO(docx_bytes))
    partes = [p.text for p in doc.paragraphs if p.text.strip()]
    for table in doc.tables:
        for row in table.rows:
            celdas = [c.text.strip() for c in row.cells]
            if any(celdas):
                partes.append(' | '.join(celdas))
    return '\n'.join(partes)


class AccountPaymentOrderQuoteImportWizard(models.TransientModel):
    _name = 'account.payment.order.quote.import.wizard'
    _description = 'Cargar Cotización de Proveedor con IA'

    order_id = fields.Many2one(
        'account.payment.order', string='Orden de Pago', required=True,
        default=lambda self: self.env.context.get('default_order_id'))
    state = fields.Selection([
        ('upload', 'Cargar Archivo'),
        ('review', 'Revisar Datos'),
    ], default='upload', required=True)
    quote_file = fields.Binary(string='Archivo de Cotización')
    quote_filename = fields.Char(string='Nombre de Archivo')
    proveedor_extraido = fields.Char(string='Proveedor Sugerido (IA)', readonly=True)
    proveedor_catalogo_id = fields.Many2one(
        'construtec.materials.vendor.mirror', string='Proveedor (Catálogo)', readonly=True,
        help='Resuelto automáticamente contra el Catálogo de Proveedores - vinculado si ya '
             'existía, creado (marcado "Pendiente de Verificar") si no.')
    fecha_extraida = fields.Date(
        string='Fecha de la Cotización (IA)', readonly=True,
        help='Solo informativa - no reemplaza la Fecha de la Orden (que ya tiene su propio '
             'valor por defecto). Cópiala a mano en el encabezado si aplica.')
    line_ids = fields.One2many(
        'account.payment.order.quote.import.wizard.line', 'wizard_id', string='Líneas Extraídas')

    def action_extraer(self):
        """Envía el archivo a la IA y pasa a la pantalla de revisión. Deliberadamente todo en
        una sola transacción (decisión explícita del usuario) - si la extracción falla, no se
        guarda ni el adjunto ni ninguna línea; el jefe de técnicos simplemente vuelve a subir el
        mismo archivo, que ya tiene en su computadora/teléfono."""
        self.ensure_one()
        if not self.quote_file:
            raise UserError(self.env._('Selecciona un archivo primero.'))
        api_key = self.order_id.company_id.anthropic_api_key
        if not api_key:
            raise UserError(self.env._(
                'No se configuró la Anthropic API Key. Ajustes > Facturación > '
                '"IA para Cotizaciones (Anthropic)".'))

        filename = (self.quote_filename or '').lower()
        file_bytes = base64.b64decode(self.quote_file)

        try:
            if filename.endswith('.pdf'):
                result = extract_quote_from_pdf(api_key, file_bytes)
            elif filename.endswith('.docx'):
                texto = _extraer_texto_docx(file_bytes)
                result = extract_quote_from_text(api_key, texto)
            else:
                ext = next((e for e in _IMAGE_MIME_BY_EXT if filename.endswith(e)), None)
                if not ext:
                    raise UserError(self.env._(
                        'Formato no soportado. Use PDF, imagen (PNG/JPG/WEBP) o Word (.docx).'))
                result = extract_quote_from_image(api_key, file_bytes, _IMAGE_MIME_BY_EXT[ext])
        except QuoteExtractionError as exc:
            raise UserError(str(exc)) from exc
        except ImportError as exc:
            raise UserError(self.env._(
                'Falta instalar la librería python-docx en el servidor: %(error)s',
                error=str(exc))) from exc

        self.env['ir.attachment'].create({
            'name': self.quote_filename,
            'datas': self.quote_file,
            'res_model': 'account.payment.order',
            'res_id': self.order_id.id,
        })

        proveedor, proveedor_creado = self._resolver_proveedor(result['proveedor'])
        self.proveedor_extraido = proveedor.name if proveedor else (result['proveedor'] or False)
        self.proveedor_catalogo_id = proveedor.id if proveedor else False
        self.fecha_extraida = result['fecha']

        materiales_conocidos = materiales_nuevos = 0
        line_commands = [(5, 0, 0)]
        for linea in result['lineas']:
            material, material_creado = self._resolver_material(
                proveedor, linea['descripcion'], linea['precio_unitario'])
            if material:
                if material_creado:
                    materiales_nuevos += 1
                else:
                    materiales_conocidos += 1
            line_commands.append((0, 0, {
                'description': linea['descripcion'],
                'qty': linea['cantidad'] or 1,
                'uom_name': linea['unidad'],
                'estimated_price': linea['precio_unitario'],
                'catalogo_id': material.id if material else False,
            }))
        self.line_ids = line_commands

        if proveedor:
            proveedor_estado = self.env._('nuevo') if proveedor_creado else self.env._('conocido')
        else:
            proveedor_estado = self.env._('no detectado')
        self.order_id.message_post(body=self.env._(
            'Cotización procesada por IA (%(filename)s): %(count)s línea(s) detectada(s), '
            'fecha: %(fecha)s. Proveedor: %(proveedor)s (%(proveedor_estado)s). '
            'Materiales: %(conocidos)s ya conocido(s), %(nuevos)s nuevo(s) (marcados '
            '"Pendiente de Verificar" en el catálogo).',
            filename=self.quote_filename, count=len(result['lineas']),
            fecha=result['fecha'] or '(no detectada)',
            proveedor=proveedor.name if proveedor else (result['proveedor'] or '(no detectado)'),
            proveedor_estado=proveedor_estado,
            conocidos=materiales_conocidos, nuevos=materiales_nuevos))
        self.state = 'review'
        return self._reload_action()

    def _resolver_proveedor(self, nombre_extraido):
        """Busca el proveedor extraído en el Catálogo de Proveedores (construtec.materials.
        vendor.mirror) por nombre normalizado - si no existe, lo crea marcado
        `pendiente_verificar` (nunca tuvo un Documento SAT real de origen, `origin_id` vacío).
        Sin proveedor detectado (`nombre_extraido` vacío), no resuelve/crea nada - devuelve un
        recordset vacío, y `_resolver_material()` tampoco intenta vincular materiales sin un
        proveedor real detrás. `sudo()` deliberado (llamador de confianza, ver el resto de este
        módulo) - `base.group_user` solo tiene lectura sobre este catálogo.

        Devuelve (record, created) - `created=True` si se creó una entrada nueva."""
        nombre_extraido = (nombre_extraido or '').strip()
        Vendor = self.env['construtec.materials.vendor.mirror'].sudo()
        if not nombre_extraido:
            return Vendor.browse(), False
        normalizado = _normalizar(nombre_extraido)
        existente = Vendor.search([]).filtered(lambda v: _normalizar(v.name) == normalizado)
        if existente:
            return existente[:1], False
        return Vendor.create({'name': nombre_extraido, 'pendiente_verificar': True}), True

    def _resolver_material(self, proveedor, descripcion, precio_unitario):
        """Busca la línea extraída en el Catálogo de Materiales (construtec.materials.catalog.
        mirror), acotado a `bien_o_servicio='B'` (mismo filtro que ya usa `catalogo_id` en
        `account.payment.order.material.line`) Y al proveedor ya resuelto - evita que
        "Cemento" de un proveedor se confunda con el de otro. Sin proveedor resuelto, no
        intenta nada (devuelve vacío) - vincular un material a un catálogo sin saber de qué
        proveedor es sería inventar un dato. Si no existe, lo crea marcado
        `pendiente_verificar`, sin `origin_id`, con el precio extraído como referencia inicial.

        Devuelve (record, created)."""
        descripcion = (descripcion or '').strip()
        Catalog = self.env['construtec.materials.catalog.mirror'].sudo()
        if not descripcion or not proveedor:
            return Catalog.browse(), False
        normalizado = _normalizar(descripcion)
        candidatos = Catalog.search([
            ('bien_o_servicio', '=', 'B'),
            ('partner_name', '=', proveedor.name),
        ])
        existente = candidatos.filtered(lambda c: _normalizar(c.name) == normalizado)
        if existente:
            return existente[:1], False
        nuevo = Catalog.create({
            'name': descripcion,
            'partner_name': proveedor.name,
            'partner_vat': proveedor.nit or False,
            'bien_o_servicio': 'B',
            'precio_referencia': precio_unitario or 0,
            'pendiente_verificar': True,
        })
        return nuevo, True

    def action_volver(self):
        self.ensure_one()
        self.line_ids = [(5, 0, 0)]
        self.proveedor_extraido = False
        self.proveedor_catalogo_id = False
        self.fecha_extraida = False
        self.state = 'upload'
        return self._reload_action()

    def action_aplicar(self):
        """Reemplaza TODA la información de materiales/proveedor ya cargada en la Orden -
        decisión explícita del usuario ("si se carga otra cotización... que reemplace toda la
        información que estaba ya colocada"): si el jefe de técnicos vuelve a usar "Cargar
        Cotización" sobre la misma Orden (ej. subió el archivo equivocado, o llegó una versión
        corregida de la cotización), la corrida anterior no debe quedar mezclada con la nueva -
        se borran todas las líneas de materiales existentes antes de crear las nuevas, y el
        proveedor sugerido se sobreescribe siempre (ya no "solo si está vacío"). Seguro porque
        esto solo corre con `order_id.state == 'borrador'` (verificado abajo) - antes de ese
        punto no existe ninguna Orden de Compra/conciliación que dependa de estas líneas."""
        self.ensure_one()
        if self.order_id.state != 'borrador':
            raise UserError(self.env._(
                'La Orden ya no está en Borrador - no se pueden agregar líneas.'))
        seleccionadas = self.line_ids.filtered(lambda line: line.incluir and line.description)
        if not seleccionadas:
            raise UserError(self.env._('Selecciona al menos una línea para aplicar.'))

        lineas_previas = len(self.order_id.material_line_ids)
        self.order_id.material_line_ids.unlink()

        Line = self.env['account.payment.order.material.line']
        for line in seleccionadas:
            Line.create({
                'order_id': self.order_id.id,
                'description': line.description,
                'qty': line.qty,
                'uom_name': line.uom_name,
                'estimated_price': line.estimated_price,
                'vendor_name': self.proveedor_extraido or False,
                'catalogo_id': line.catalogo_id.id if line.catalogo_id else False,
            })
        # `proveedor_materiales_name` (el texto que viaja a Enterprise) se deriva solo - ver
        # `account.payment.order._sync_proveedor_materiales_name()`, disparada por este mismo
        # write() al traer `proveedor_materiales_catalogo_id` en vals.
        self.order_id.proveedor_materiales_catalogo_id = self.proveedor_catalogo_id.id

        if lineas_previas:
            self.order_id.message_post(body=self.env._(
                '%(count)s línea(s) de materiales aplicada(s) desde una nueva cotización '
                'cargada con IA - se reemplazaron %(previas)s línea(s) que ya existían.',
                count=len(seleccionadas), previas=lineas_previas))
        else:
            self.order_id.message_post(body=self.env._(
                '%(count)s línea(s) de materiales aplicada(s) desde una cotización cargada '
                'con IA.', count=len(seleccionadas)))
        return {'type': 'ir.actions.act_window_close'}

    def _reload_action(self):
        return {
            'type': 'ir.actions.act_window',
            'res_model': self._name,
            'res_id': self.id,
            'view_mode': 'form',
            'target': 'new',
        }


class AccountPaymentOrderQuoteImportWizardLine(models.TransientModel):
    _name = 'account.payment.order.quote.import.wizard.line'
    _description = 'Línea Extraída de una Cotización (revisión antes de aplicar)'

    wizard_id = fields.Many2one(
        'account.payment.order.quote.import.wizard', required=True, ondelete='cascade')
    incluir = fields.Boolean(string='Incluir', default=True)
    description = fields.Char(string='Descripción')
    qty = fields.Float(string='Cantidad', default=1)
    uom_name = fields.Char(string='Unidad de Medida')
    estimated_price = fields.Float(string='Precio Estimado')
    subtotal = fields.Float(string='Subtotal', compute='_compute_subtotal')
    catalogo_id = fields.Many2one(
        'construtec.materials.catalog.mirror', string='Producto SAT', readonly=True,
        help='Resuelto automáticamente contra el Catálogo de Materiales - vinculado si ya '
             'existía para este proveedor, creado (marcado "Pendiente de Verificar") si no.')

    @api.depends('qty', 'estimated_price')
    def _compute_subtotal(self):
        for line in self:
            line.subtotal = line.qty * line.estimated_price
